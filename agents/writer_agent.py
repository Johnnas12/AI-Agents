from agent import Agent
import google.generativeai as genai
from docx import Document

class WriterAgent(Agent):
    def run(self, research_text):
        print(f"[{self.name}] Drafting technical report ...")

        document = self.draft_report(research_text)
        return document
    
    def draft_report(self, research_text):

        genai.configure(api_key = self.api_key)
        prompt = f"""You are a technical writer. Using the following research text, organize it into a clear, properly structured, polished technical report in markdown format. Include a title, table of contents, headings, and a conclusion.
        Research Text:
        {research_text}
        """
        model = genai.GenerativeModel(self.model)  
        response = model.generate_content(prompt)

        report_text = response.candidates[0].content.parts[0].text
        return report_text
    
    def revise_document(self, document_text, fact_check_feedback):
        print(f"[{self.name}] Revising document based on fact-check feedback ...")

        genai.configure(api_key = self.api_key)
        model = genai.GenerativeModel(self.model)
 
        prompt = f"""You are an AI technical writer.

            Here is your original document:
            {document_text}

            And here is a list of fact-check comments you need to address:
            {fact_check_feedback}

            Revise the document accordingly, making sure to fix any inaccuracies. Keep formatting clean. Return only the updated document text.
            """
        response = model.generate_content(prompt)
        revised_document = response.candidates[0].content.parts[0].text

        return revised_document
    
    def save_to_docx(self, document_text, filename):
        print(f"[{self.name}] Saving document to {filename} ...")

        doc = Document()
        doc.add_heading('Technical Report', level=1)
        doc.add_paragraph(document_text)
        doc.save(filename)
        print(f"[{self.name}] Document saved successfully.")


